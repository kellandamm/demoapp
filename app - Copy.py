import os
import uuid
import json
import pickle
import shutil
import logging
import tempfile
import openai
from datetime import datetime, timedelta
from io import BytesIO
from flask_sqlalchemy import SQLAlchemy
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_session import Session
import requests
from azure.core.credentials import AzureKeyCredential
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    SearchIndex,
    SearchField,
    SearchFieldDataType,
    SimpleField,
    SearchableField,
    VectorSearch,
    HnswAlgorithmConfiguration,
    VectorSearchProfile,
    VectorSearchAlgorithmKind,
    VectorSearchAlgorithmMetric,
)
import PyPDF2
import docx
import base64
import hashlib
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    jsonify,
    send_file,
    session,
)

from flask_login import (
    LoginManager,
    UserMixin,
    login_user,
    logout_user,
    login_required,
    current_user,
)

load_dotenv()

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Get database configuration from environment variables
db_user = os.environ.get("DB-USER")
db_password = os.environ.get("DB-PASSWORD")
db_host = os.environ.get("DB-HOST")
db_port = os.environ.get("DB-PORT")
db_name = os.environ.get("DB-NAME")

app = Flask(__name__)
app.config["BASE_DIR"] = "user_data"
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024
app.config["SQLALCHEMY_DATABASE_URI"] = (
    f"postgresql://{db_user}:{db_password}@{db_host}:{db_port}/{db_name}"
)
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", str(uuid.uuid4()))
app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_PERMANENT"] = True
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=7)
app.config["SESSION_USE_SIGNER"] = True
app.config["SESSION_FILE_THRESHOLD"] = 500
app.config["SESSION_FILE_DIR"] = "/tmp/flask_session"

# Azure OpenAI Configuration
app.config["AZURE_OPENAI_ENDPOINT"] = os.environ.get("AZURE_OPENAI_ENDPOINT")
app.config["AZURE_OPENAI_KEY"] = os.environ.get("AZURE_OPENAI_KEY")
app.config["AZURE_OPENAI_API_VERSION"] = os.environ.get(
    "AZURE_OPENAI_API_VERSION", "2023-05-15"
)
app.config["AZURE_OPENAI_DEPLOYMENT"] = os.environ.get("AZURE_OPENAI_DEPLOYMENT")
app.config["AZURE_OPENAI_EMBEDDING_DEPLOYMENT"] = os.environ.get(
    "AZURE_OPENAI_EMBEDDING_DEPLOYMENT"
)

# Azure AI Search Configuration
app.config["AZURE_SEARCH_ENDPOINT"] = os.environ.get("AZURE_SEARCH_ENDPOINT")
app.config["AZURE_SEARCH_KEY"] = os.environ.get("AZURE_SEARCH_KEY")
app.config["AZURE_SEARCH_INDEX_NAME"] = os.environ.get("AZURE_SEARCH_INDEX_NAME")
app.logger.setLevel(logging.DEBUG)


# Add template filter for timestamp formatting
@app.template_filter("timestamp_to_date")
def timestamp_to_date(timestamp):
    """Convert UNIX timestamp to formatted date string"""
    if not timestamp:
        return ""
    return datetime.fromtimestamp(timestamp).strftime("%Y-%m-%d %H:%M:%S")


login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"  # Specify the login view
login_manager.session_protection = "strong"
os.makedirs(app.config["SESSION_FILE_DIR"], exist_ok=True)

db = SQLAlchemy(app)
Session(app)


class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(255), nullable=False)
    content_type = db.Column(db.String(100), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey("user_account.id"), nullable=False)
    is_indexed = db.Column(db.Boolean, default=False)
    chunk_count = db.Column(db.Integer, default=0)

    # Define a relationship to User model
    user = db.relationship("User", backref=db.backref("documents", lazy=True))

    def __repr__(self):
        return f"<Document {self.filename}>"


class DocumentProcessor:
    def __init__(self, app):
        self.app = app
        self.openai_endpoint = app.config.get("AZURE_OPENAI_ENDPOINT")
        self.openai_key = app.config.get("AZURE_OPENAI_KEY")
        self.api_version = app.config.get("AZURE_OPENAI_API_VERSION")
        self.embedding_deployment = app.config.get("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")

        # Initialize OpenAI client if credentials provided
        self.initialized = False
        if self.openai_endpoint and self.openai_key and self.embedding_deployment:
            try:
                self.client = openai.AzureOpenAI(
                    api_key=self.openai_key,
                    api_version=self.api_version,
                    azure_endpoint=self.openai_endpoint,
                )
                self.initialized = True
                app.logger.info("Document processor initialized successfully")
            except Exception as e:
                app.logger.error(f"Failed to initialize document processor: {str(e)}")

    def extract_text(self, file_path, content_type):
        """Extract text from various file types"""
        try:
            text = ""

            # Handle different file types
            if content_type == "application/pdf":
                # Extract text from PDF
                with open(file_path, "rb") as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n\n"

            elif (
                content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                # Extract text from DOCX
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"

            elif content_type == "text/plain":
                # Handle text files
                with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                    text = file.read()

            else:
                # For unsupported formats, return empty string
                self.app.logger.warning(f"Unsupported content type: {content_type}")
                return ""

            return text
        except Exception as e:
            self.app.logger.error(f"Error extracting text from file: {str(e)}")
            return ""

    def chunk_text(self, text, chunk_size=1000, overlap=100):
        """Split text into chunks with overlap"""
        if not text:
            return []

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + chunk_size
            if end >= text_length:
                # Last chunk
                chunk_text = text[start:text_length]
                chunks.append(chunk_text)
                break

            # Find a good breaking point (newline or space)
            # Look for newline first
            cutoff = text.rfind("\n", start, end)
            if cutoff == -1 or cutoff <= start:
                # No good newline found, look for space
                cutoff = text.rfind(" ", start, end)
                if cutoff == -1 or cutoff <= start:
                    # No good space found, just cut at the maximum length
                    cutoff = end

            chunks.append(text[start:cutoff])
            start = cutoff - overlap if cutoff > overlap else cutoff

        return chunks

    def get_embeddings(self, text_chunks):
        """Get embeddings for text chunks"""
        if not self.initialized:
            self.app.logger.warning("Document processor not initialized")
            return []

        try:
            chunk_with_embeddings = []

            for chunk in text_chunks:
                try:
                    # Skip empty chunks
                    if not chunk.strip():
                        continue

                    response = self.client.embeddings.create(
                        input=chunk, model=self.embedding_deployment
                    )

                    embedding = response.data[0].embedding
                    chunk_with_embeddings.append(
                        {"content": chunk, "embedding": embedding}
                    )

                except Exception as e:
                    self.app.logger.error(f"Error getting embedding: {str(e)}")
                    # Continue with other chunks

            return chunk_with_embeddings
        except Exception as e:
            self.app.logger.error(f"Error processing document chunks: {str(e)}")
            return []

    def process_file(self, file_path, content_type):
        """Process a file and return chunks with embeddings"""
        # Extract text from file
        text = self.extract_text(file_path, content_type)
        if not text:
            return []

        # Split text into chunks
        text_chunks = self.chunk_text(text)

        # Get embeddings for chunks
        return self.get_embeddings(text_chunks)


class AISearchManager:
    def __init__(self, app):
        self.app = app
        self.endpoint = app.config.get("AZURE_SEARCH_ENDPOINT")
        self.key = app.config.get("AZURE_SEARCH_KEY")
        self.index_name = app.config.get("AZURE_SEARCH_INDEX_NAME")
        self.dimension = 1536  # OpenAI embedding dimension

        # Initialize clients if credentials provided
        self.initialized = False
        if self.endpoint and self.key:
            try:
                self.credential = AzureKeyCredential(self.key)
                self.index_client = SearchIndexClient(
                    endpoint=self.endpoint, credential=self.credential
                )
                self.search_client = SearchClient(
                    endpoint=self.endpoint,
                    index_name=self.index_name,
                    credential=self.credential,
                )
                self._ensure_index_exists()
                self.initialized = True
                app.logger.info("Azure AI Search initialized successfully")
            except Exception as e:
                app.logger.error(f"Failed to initialize Azure AI Search: {str(e)}")

    def _ensure_index_exists(self):
        """Create the search index if it doesn't exist"""
        try:
            # Check if index exists
            if self.index_name in [
                index.name for index in self.index_client.list_indexes()
            ]:
                self.app.logger.info(f"Search index '{self.index_name}' already exists")
                return True

            # Create a new index
            vector_search = VectorSearch(
                algorithms=[
                    HnswAlgorithmConfiguration(
                        name="hnsw-config",
                        parameters={
                            "m": 4,
                            "efConstruction": 400,
                            "efSearch": 500,
                            "metric": "cosine",
                        },
                    )
                ],
                profiles=[
                    VectorSearchProfile(
                        name="vector-profile",
                        algorithm_configuration_name="hnsw-config",
                        vectorizer="none",
                    )
                ],
            )

            fields = [
                SimpleField(name="id", type=SearchFieldDataType.String, key=True),
                SimpleField(
                    name="user_id", type=SearchFieldDataType.String, filterable=True
                ),
                SimpleField(
                    name="document_id", type=SearchFieldDataType.String, filterable=True
                ),
                SimpleField(
                    name="filename",
                    type=SearchFieldDataType.String,
                    filterable=True,
                    sortable=True,
                ),
                SimpleField(name="chunk_id", type=SearchFieldDataType.Int32),
                SearchableField(name="content", type=SearchFieldDataType.String),
                SearchableField(
                    name="embedding",
                    type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                    vector_search_dimensions=self.dimension,
                    vector_search_profile_name="vector-profile",
                ),
            ]

            index = SearchIndex(
                name=self.index_name, fields=fields, vector_search=vector_search
            )
            self.index_client.create_index(index)
            self.app.logger.info(f"Created search index '{self.index_name}'")
            return True
        except Exception as e:
            self.app.logger.error(f"Error creating search index: {str(e)}")
            return False

    def upload_document_chunks(self, user_id, document_id, filename, chunks):
        """Upload document chunks to the search index"""
        if not self.initialized:
            self.app.logger.warning("Azure AI Search not initialized")
            return False

        try:
            batch_documents = []

            for i, chunk in enumerate(chunks):
                content = chunk["content"]
                embedding = chunk["embedding"]

                doc = {
                    "id": f"{user_id}-{document_id}-{i}",
                    "user_id": str(user_id),
                    "document_id": str(document_id),
                    "filename": filename,
                    "chunk_id": i,
                    "content": content,
                    "embedding": embedding,
                }
                batch_documents.append(doc)

            # Upload in batches of 100
            batch_size = 100
            for i in range(0, len(batch_documents), batch_size):
                batch = batch_documents[i : i + batch_size]
                self.search_client.upload_documents(documents=batch)

            self.app.logger.info(
                f"Uploaded {len(chunks)} chunks for document {filename}"
            )
            return True
        except Exception as e:
            self.app.logger.error(f"Error uploading document chunks: {str(e)}")
            return False

    def search_documents(self, user_id, query_embedding, top_k=5):
        """Search documents by vector similarity"""
        if not self.initialized:
            self.app.logger.warning("Azure AI Search not initialized")
            return []

        try:
            # Build filter for user's documents only
            filter_expr = f"user_id eq '{user_id}'"

            # Perform vector search
            results = self.search_client.search(
                search_text=None,
                vector={"vector": query_embedding, "k": top_k, "fields": "embedding"},
                filter=filter_expr,
                select=["id", "filename", "content", "chunk_id"],
                top=top_k,
            )

            # Process results
            search_results = []
            for result in results:
                search_results.append(
                    {
                        "id": result["id"],
                        "filename": result["filename"],
                        "content": result["content"],
                        "chunk_id": result["chunk_id"],
                        "score": result["@search.score"],
                    }
                )

            return search_results
        except Exception as e:
            self.app.logger.error(f"Error searching documents: {str(e)}")
            return []

    def delete_document(self, user_id, document_id):
        """Delete all chunks for a document"""
        if not self.initialized:
            self.app.logger.warning("Azure AI Search not initialized")
            return False

        try:
            # Find all chunks for this document
            filter_expr = f"user_id eq '{user_id}' and document_id eq '{document_id}'"
            results = self.search_client.search(
                search_text="*",
                filter=filter_expr,
                select=["id"],
                top=1000,  # Adjust if needed for large documents
            )

            # Delete each document
            ids_to_delete = [result["id"] for result in results]
            if ids_to_delete:
                batch_actions = [
                    {"@search.action": "delete", "id": doc_id}
                    for doc_id in ids_to_delete
                ]
                self.search_client.index_documents(batch=batch_actions)
                self.app.logger.info(
                    f"Deleted {len(ids_to_delete)} chunks for document {document_id}"
                )
            return True
        except Exception as e:
            self.app.logger.error(f"Error deleting document chunks: {str(e)}")
            return False  # Class definitions should be placed before the instantiation


# Initialize document processor and search manager
document_processor = DocumentProcessor(app)
search_manager = AISearchManager(
    app
)  # Class definitions should be placed before the instantiation


@app.route("/api/chat/documents/query", methods=["POST"])
@login_required
def document_query():
    """Query user's documents with natural language"""
    try:
        data = request.json
        query = data.get("query")
        document_ids = data.get(
            "document_ids", []
        )  # Optional filter by specific documents

        if not query:
            return jsonify({"error": "No query provided"}), 400

        if not document_processor.initialized or not search_manager.initialized:
            return jsonify(
                {"error": "Document search not available. Check server configuration."}
            ), 500

        # Get embedding for the query
        response = document_processor.client.embeddings.create(
            input=query, model=document_processor.embedding_deployment
        )
        query_embedding = response.data[0].embedding

        # Search for relevant content
        search_results = search_manager.search_documents(
            current_user.id, query_embedding, top_k=5
        )

        if not search_results:
            return jsonify(
                {
                    "answer": "I couldn't find any relevant information in your documents.",
                    "sources": [],
                }
            )

        # Format context from search results
        context = ""
        sources = []

        for result in search_results:
            context += f"\nPassage: {result['content']}\n"
            sources.append(
                {
                    "filename": result["filename"],
                    "content": result["content"][:200] + "...",  # Preview
                    "score": result["score"],
                }
            )

        # Build prompt with context
        system_message = """You are a helpful assistant that answers questions based on the user's documents. 
        Use ONLY the provided document passages to answer the question. If the answer cannot be found in the 
        passages, say "I couldn't find information about this in your documents." and suggest what kinds of 
        documents might have this information."""

        user_message = f"""
        Question: {query}
        
        Here are relevant passages from my documents:
        {context}
        
        Please answer my question based on these passages only.
        """

        # Get response from Azure OpenAI
        response = get_azure_openai_response(
            [{"role": "user", "content": user_message}], system_message
        )

        return jsonify({"answer": response, "sources": sources})

    except Exception as e:
        app.logger.error(f"Error querying documents: {str(e)}")
        return jsonify({"error": f"Error processing your query: {str(e)}"}), 500


@login_required
def reindex_document(document_id):
    """Re-index a document"""
    try:
        document = Document.query.get_or_404(document_id)

        # Verify ownership
        if document.user_id != current_user.id:
            return jsonify(
                {"success": False, "error": "You don't own this document"}
            ), 403

        # Delete existing index if any
        if document.is_indexed and search_manager.initialized:
            search_manager.delete_document(document.user_id, document.id)

        # Reset indexing status
        document.is_indexed = False
        document.chunk_count = 0
        db.session.commit()

        # Reindex the document
        success = index_document(document.id)

        return jsonify(
            {
                "success": success,
                "is_indexed": document.is_indexed,
                "chunk_count": document.chunk_count,
            }
        )
    except Exception as e:
        app.logger.error(f"Error reindexing document: {str(e)}")
        return jsonify({"success": False, "error": str(e)}) @ app.route(
            "/api/documents/<int:document_id>/delete", methods=["POST"]
        )


@login_required
def delete_document_api(document_id):
    """Delete a document and its index"""
    try:
        document = Document.query.get_or_404(document_id)

        # Verify ownership
        if document.user_id != current_user.id:
            return jsonify(
                {"success": False, "error": "You don't own this document"}
            ), 403

        # Delete from search index first
        if document.is_indexed and search_manager.initialized:
            search_manager.delete_document(document.user_id, document.id)

        # Delete file from disk
        if os.path.exists(document.file_path):
            os.remove(document.file_path)

        # Delete database record
        db.session.delete(document)
        db.session.commit()

        return jsonify({"success": True})
    except Exception as e:
        app.logger.error(f"Error deleting document: {str(e)}")
        return jsonify({"success": False, "error": str(e)})


def index_document(document_id):
    """Process and index a document"""
    try:
        # Get document from database
        document = Document.query.get(document_id)
        if not document:
            app.logger.error(f"Document {document_id} not found")
            return False

        # Check if document processor and search manager are initialized
        if not document_processor.initialized or not search_manager.initialized:
            app.logger.error("Document processor or search manager not initialized")
            document.is_indexed = False
            db.session.commit()
            return False

        # Process the document
        chunks_with_embeddings = document_processor.process_file(
            document.file_path, document.content_type
        )
        if not chunks_with_embeddings:
            app.logger.error(f"Failed to process document {document.filename}")
            document.is_indexed = False
            db.session.commit()
            return False

        # Upload chunks to search index
        success = search_manager.upload_document_chunks(
            document.user_id, document.id, document.filename, chunks_with_embeddings
        )

        if success:
            # Update document status
            document.is_indexed = True
            document.chunk_count = len(chunks_with_embeddings)
            db.session.commit()
            app.logger.info(
                f"Successfully indexed document {document.filename} with {len(chunks_with_embeddings)} chunks"
            )
            return True
        else:
            document.is_indexed = False
            db.session.commit()
            return False

    except Exception as e:
        app.logger.error(f"Error indexing document: {str(e)}")
        return False


def get_azure_openai_response(messages, system_message=None):
    """Get a response from Azure OpenAI with robust error handling"""
    try:
        if (
            not app.config["AZURE_OPENAI_KEY"]
            or not app.config["AZURE_OPENAI_ENDPOINT"]
        ):
            return "Azure OpenAI is not configured. Please set the AZURE_OPENAI_KEY and AZURE_OPENAI_ENDPOINT environment variables."

        client = openai.AzureOpenAI(
            api_key=app.config["AZURE_OPENAI_KEY"],
            api_version=app.config["AZURE_OPENAI_API_VERSION"],
            azure_endpoint=app.config["AZURE_OPENAI_ENDPOINT"],
        )

        api_messages = []
        if system_message:
            api_messages.append({"role": "system", "content": system_message})

        # Convert our messages to the format expected by the API
        for msg in messages:
            if msg["role"] in ["user", "assistant", "system"]:
                api_messages.append({"role": msg["role"], "content": msg["content"]})

        response = client.chat.completions.create(
            model=app.config["AZURE_OPENAI_DEPLOYMENT"],
            messages=api_messages,
            temperature=0.7,
            max_tokens=800,
        )

        return response.choices[0].message.content
    except Exception as e:
        app.logger.error(f"Azure OpenAI error: {str(e)}")
        return f"I encountered an error connecting to the AI service: {str(e)}. Please try again later."


# Initialize document processor and search manager
document_processor = DocumentProcessor(app)
search_manager = AISearchManager(app)


# User model
class User(UserMixin, db.Model):
    __tablename__ = "user_account"
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    is_admin = db.Column(db.Boolean, default=False)
    is_approved = db.Column(db.Boolean, default=False)  # Approval status
    approval_date = db.Column(db.DateTime, nullable=True)  # When approved
    registration_date = db.Column(db.DateTime, default=datetime.utcnow)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    @property
    def is_active(self):
        # Only allow login if approved (override UserMixin)
        return self.is_approved


class Environment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=True)
    packages = db.Column(db.Text, nullable=False)  # JSON string of packages
    is_public = db.Column(db.Boolean, default=False)
    user_id = db.Column(db.Integer, db.ForeignKey("user_account.id"), nullable=False)

    def get_path(self):
        return os.path.join(
            app.config["BASE_DIR"], str(self.user_id), "environments", self.name
        )

    def get_packages_list(self):
        try:
            return json.loads(self.packages)
        except:
            return []

    def get_variable_store_path(self):
        return os.path.join(
            app.config["BASE_DIR"],
            str(self.user_id),
            "variables",
            f"{os.path.splitext(self.filename)[0]}.pkl",
        )


class FileStorage:
    """Simple file storage that uses local filesystem"""

    def __init__(self, app):
        self.app = app

    def _ensure_user_dirs(self, user_id):
        """Ensure user directories exist for local storage"""
        base_dir = os.path.join(self.app.config["BASE_DIR"], str(user_id))
        uploads_dir = os.path.join(base_dir, "uploads")
        os.makedirs(uploads_dir, exist_ok=True)
        return uploads_dir

    def upload_file(self, file_stream, filename, user_id):
        """Upload a file to storage"""
        secure_name = secure_filename(filename)
        uploads_dir = self._ensure_user_dirs(user_id)
        filepath = os.path.join(uploads_dir, secure_name)
        file_stream.save(filepath)
        return filepath

    def download_file(self, filename, user_id):
        """Download a file from storage"""
        secure_name = secure_filename(filename)
        uploads_dir = self._ensure_user_dirs(user_id)
        filepath = os.path.join(uploads_dir, secure_name)
        if os.path.exists(filepath):
            with open(filepath, "rb") as f:
                return f.read()
        return None

    def delete_file(self, filename, user_id):
        """Delete a file from storage"""
        secure_name = secure_filename(filename)
        uploads_dir = self._ensure_user_dirs(user_id)
        filepath = os.path.join(uploads_dir, secure_name)
        if os.path.exists(filepath):
            os.remove(filepath)

    def list_files(self, user_id):
        """List all files for a user"""
        uploads_dir = self._ensure_user_dirs(user_id)
        files = []

        for filename in os.listdir(uploads_dir):
            file_path = os.path.join(uploads_dir, filename)
            if os.path.isfile(file_path):
                files.append(
                    {
                        "name": filename,
                        "size": os.path.getsize(file_path),
                        "modified": os.path.getmtime(file_path),
                        "filepath": file_path,
                    }
                )

        return files


file_storage = FileStorage(app)


@login_manager.user_loader
def load_user(user_id):
    try:
        # Convert to int since user_id is stored as string in session
        return User.query.get(int(user_id))
    except (ValueError, TypeError):
        # Handle invalid user_id values
        return None


# Ensure user directories exist
def ensure_user_dirs(user_id):
    base_dir = os.path.join(app.config["BASE_DIR"], str(user_id))
    uploads_dir = os.path.join(base_dir, "uploads")
    variables_dir = os.path.join(base_dir, "variables")

    for directory in [base_dir, uploads_dir, variables_dir]:
        os.makedirs(directory, exist_ok=True)

    return uploads_dir, variables_dir


@app.before_request
def check_session_consistency():
    """Ensure user session is consistent"""
    if current_user.is_authenticated:
        # Check if session has user_id and it matches current_user.id
        if session.get("user_id") != current_user.id:
            app.logger.warning(
                f"Session inconsistency detected. Logged in as {current_user.id} but session has {session.get('user_id')}"
            )
            logout_user()
            session.clear()
            flash("Your session has expired. Please log in again.")
            return redirect(url_for("login"))


@app.route("/documents")
@login_required
def document_list():
    """Document management page"""
    documents = (
        Document.query.filter_by(user_id=current_user.id)
        .order_by(Document.upload_date.desc())
        .all()
    )
    return render_template("documents.html", documents=documents)


@app.route("/")
def index():
    app.logger.debug(
        f"Accessing index route. current_user.is_authenticated: {current_user.is_authenticated}, current_user.id: {current_user.id if current_user.is_authenticated else None}, Session: {session}"
    )
    # Don't check current_user.is_authenticated which might be unreliable
    # Instead, check if user_id is in session
    if session.get("user_id"):
        # Get recent conversations for display on the index page
        chat_history_path = os.path.join(
            app.config["BASE_DIR"], str(session.get("user_id")), "chat_history"
        )
        recent_conversations = []

        try:
            if os.path.exists(chat_history_path):
                files = sorted(
                    [f for f in os.listdir(chat_history_path) if f.endswith(".json")],
                    key=lambda x: os.path.getmtime(os.path.join(chat_history_path, x)),
                    reverse=True,
                )[:5]  # Get 5 most recent files

                for filename in files:
                    filepath = os.path.join(chat_history_path, filename)
                    try:
                        with open(filepath, "r") as f:
                            data = json.load(f)
                            conversation_id = os.path.splitext(filename)[0]
                            title = data[0]["content"][:30] + (
                                "..." if len(data[0]["content"]) > 30 else ""
                            )
                            timestamp = datetime.fromtimestamp(
                                os.path.getmtime(filepath)
                            )
                            recent_conversations.append(
                                {
                                    "id": conversation_id,
                                    "title": title,
                                    "timestamp": timestamp.strftime(
                                        "%Y-%m-%d %H:%M:%S"
                                    ),
                                }
                            )
                    except:
                        pass
        except Exception as e:
            app.logger.error(f"Error loading recent conversations: {str(e)}")

        # Get recent files
        recent_files = file_storage.list_files(session.get("user_id"))
        if len(recent_files) > 5:
            recent_files = sorted(
                recent_files, key=lambda x: x["modified"], reverse=True
            )[:5]

        return render_template(
            "index.html",
            recent_conversations=recent_conversations,
            recent_files=recent_files,
        )
    else:
        return redirect(url_for("login"))


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        email = request.form.get(
            "email", f"{username}@example.com"
        )  # Default email if not provided

        if User.query.filter_by(username=username).first():
            flash("Username already exists")
            return redirect(url_for("register"))

        # Create new user
        new_user = User(
            username=username,
            email=email,
        )
        new_user.set_password(password)

        # Auto-approve admin account
        if username == "admin":
            new_user.is_admin = True
            new_user.is_approved = True

        db.session.add(new_user)
        db.session.commit()

        # Create user directories
        ensure_user_dirs(new_user.id)

        flash("Account created successfully!")
        return redirect(url_for("login"))

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    # Check if already logged in via session (not current_user)
    if session.get("user_id"):
        return redirect(url_for("index"))

    try:
        if request.method == "POST":
            username = request.form.get("username")
            password = request.form.get("password")

            user = User.query.filter_by(username=username).first()
            if user and user.check_password(password):
                if not user.is_approved:
                    flash("Your account is pending approval.")
                    return render_template("login.html")

                # Login with Flask-Login
                login_user(user, remember=True)

                # CRITICAL: Set session data explicitly
                session["user_id"] = user.id
                session["username"] = user.username
                session.modified = True  # Ensure session is saved

                app.logger.debug(
                    f"Login successful. User ID: {user.id}, Session: {session}, Redirecting to: {url_for('index')}"
                )

                # Always redirect to index page after successful login
                return redirect(url_for("index"))
            else:
                app.logger.debug("Invalid username or password")
                flash("Invalid username or password")
                return render_template("login.html")

        return render_template("login.html")
    except Exception as e:
        app.logger.error(f"Login error: {str(e)}")
        flash(f"System error: {str(e)}")
        return render_template("login.html")


@app.route("/logout")
def logout():
    # Clear Flask-Login authentication
    logout_user()

    # Clear all session data
    session.clear()

    flash("You have been logged out.")
    return redirect(url_for("login"))


@app.route("/admin")
@login_required
def admin():
    if not current_user.is_admin:
        flash("You do not have permission to access the admin page")
        return redirect(url_for("index"))

    users = User.query.all()

    # Count pending users
    pending_count = User.query.filter_by(is_approved=False).count()

    return render_template("admin.html", users=users, pending_count=pending_count)


@app.route("/admin/pending_users")
@login_required
def pending_users():
    """Show pending user approvals"""
    if not current_user.is_admin:
        flash("You do not have permission to access this page")
        return redirect(url_for("index"))

    # Get users pending approval
    pending = (
        User.query.filter_by(is_approved=False).order_by(User.registration_date).all()
    )

    return render_template("pending_users.html", pending_users=pending)


@app.route("/admin/approve_user/<int:user_id>", methods=["POST"])
@login_required
def approve_user(user_id):
    """Approve a pending user"""
    if not current_user.is_admin:
        flash("You do not have permission to perform this action")
        return redirect(url_for("index"))

    user = User.query.get_or_404(user_id)

    # Update approval status
    user.is_approved = True
    user.approval_date = datetime.utcnow()
    db.session.commit()

    flash(f"User {user.username} has been approved")
    return redirect(url_for("pending_users"))


@app.route("/admin/reject_user/<int:user_id>", methods=["POST"])
@login_required
def reject_user(user_id):
    """Reject and delete a pending user"""
    if not current_user.is_admin:
        flash("You do not have permission to perform this action")
        return redirect(url_for("index"))

    user = User.query.get_or_404(user_id)

    # Delete user directories if they exist
    user_dir = os.path.join(app.config["BASE_DIR"], str(user.id))
    if os.path.exists(user_dir):
        shutil.rmtree(user_dir)

    # Delete user from database
    db.session.delete(user)
    db.session.commit()

    flash(f"User {user.username} has been rejected and removed")
    return redirect(url_for("pending_users"))


@app.route("/api/files")
@login_required
def list_files():
    """API to list files in user's storage"""
    files = file_storage.list_files(current_user.id)
    return jsonify({"success": True, "files": files})


@app.route("/api/files/upload", methods=["POST"])
@login_required
def api_upload_file():
    """API to upload multiple files to storage"""
    if "files" not in request.files:
        return jsonify({"success": False, "error": "No file part"})

    files = request.files.getlist("files")
    
    if not files or files[0].filename == "":
        return jsonify({"success": False, "error": "No files selected"})
    
    results = []
    
    for file in files:
        try:
            # Upload file to storage
            file_id = file_storage.upload_file(file, file.filename, current_user.id)
            
            # Create document record and index it
            new_document = Document(
                filename=secure_filename(file.filename),
                file_path=file_id,
                content_type=file.content_type,
                user_id=current_user.id,
                is_indexed=False
            )
            db.session.add(new_document)
            db.session.commit()
            
            # Start indexing
            index_document(new_document.id)
            
            results.append({
                "filename": file.filename,
                "success": True,
                "file_id": file_id
            })
        except Exception as e:
            app.logger.error(f"Upload error for {file.filename}: {str(e)}")
            results.append({
                "filename": file.filename,
                "success": False,
                "error": str(e)
            })
    
    return jsonify({
        "success": any(r["success"] for r in results),
        "results": results
    })


@app.route("/api/files/delete/<filename>", methods=["DELETE"])
@login_required
def delete_file(filename):
    """API to delete a file from storage"""
    try:
        file_storage.delete_file(filename, current_user.id)
        return jsonify({"success": True})
    except Exception as e:
        app.logger.error(f"Delete error: {str(e)}")
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/files/download/<filename>")
@login_required
def download_file(filename):
    """API to download a file from storage"""
    try:
        file_data = file_storage.download_file(filename, current_user.id)

        if file_data is None:
            flash("File not found")
            return redirect(url_for("file_upload"))

        # Create BytesIO object for sending
        file_stream = BytesIO(file_data)

        return send_file(file_stream, download_name=filename, as_attachment=True)
    except Exception as e:
        app.logger.error(f"Download error: {str(e)}")
        flash(f"Error downloading file: {str(e)}")
        return redirect(url_for("file_upload"))


@app.route("/admin/delete_user/<int:user_id>", methods=["POST"])
@login_required
def delete_user(user_id):
    """Delete a user and all their associated data"""
    # Check if current user is admin
    if not current_user.is_admin:
        flash("You do not have permission to delete users")
        return redirect(url_for("index"))

    # Prevent admin from deleting themselves
    if user_id == current_user.id:
        flash("You cannot delete your own admin account")
        return redirect(url_for("admin"))

    # Get the user to delete
    user_to_delete = User.query.get_or_404(user_id)

    # Delete user directories if they exist
    user_dir = os.path.join(app.config["BASE_DIR"], str(user_to_delete.id))
    if os.path.exists(user_dir):
        shutil.rmtree(user_dir)

    # Delete the user
    db.session.delete(user_to_delete)
    db.session.commit()

    flash(f"User {user_to_delete.username} and all associated data have been deleted")
    return redirect(url_for("admin"))


@app.route("/admin/change_password/<int:user_id>", methods=["POST"])
@login_required
def change_user_password(user_id):
    """Change a user's password (admin only)"""
    # Check if current user is admin
    if not current_user.is_admin:
        flash("You do not have permission to change user passwords")
        return redirect(url_for("index"))

    # Get the user
    user = User.query.get_or_404(user_id)

    # Get the new password from the form
    new_password = request.form.get("new_password")
    confirm_password = request.form.get("confirm_password")

    # Validate passwords
    if not new_password or len(new_password) < 4:
        flash("Password must be at least 4 characters long")
        return redirect(url_for("admin"))

    if new_password != confirm_password:
        flash("Passwords do not match")
        return redirect(url_for("admin"))

    # Update the password
    user.set_password(new_password)
    db.session.commit()

    flash(f"Password updated for user: {user.username}")
    return redirect(url_for("admin"))


@app.route("/admin/toggle_admin/<int:user_id>", methods=["POST"])
@login_required
def toggle_admin(user_id):
    """Toggle admin status for a user"""
    # Check if current user is admin
    if not current_user.is_admin:
        flash("You do not have permission to change user roles")
        return redirect(url_for("index"))

    # Prevent changing self (to avoid locking out)
    if user_id == current_user.id:
        flash("You cannot change your own admin status")
        return redirect(url_for("admin"))

    # Get the user
    user = User.query.get_or_404(user_id)

    # Toggle admin status
    user.is_admin = not user.is_admin
    db.session.commit()

    flash(
        f"User {user.username} is now {'an admin' if user.is_admin else 'a normal user'}"
    )
    return redirect(url_for("admin"))


@app.route("/profile")
@login_required
def profile():
    """User profile page"""
    return render_template("profile.html")


@app.route("/profile/update", methods=["POST"])
@login_required
def update_profile():
    """Update user profile information"""
    username = request.form.get("username")

    # Validate username
    if not username:
        flash("Username cannot be empty")
        return redirect(url_for("profile"))

    # Check if username is taken (if changed)
    if username != current_user.username:
        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            flash("Username already taken")
            return redirect(url_for("profile"))

        # Update username
        current_user.username = username

    # Add email if form has it
    email = request.form.get("email")
    if email:
        current_user.email = email

    db.session.commit()
    flash("Profile updated successfully")
    return redirect(url_for("profile"))


@app.route("/profile/change_password", methods=["POST"])
@login_required
def change_own_password():
    """Change own password"""
    current_password = request.form.get("current_password")
    new_password = request.form.get("new_password")
    confirm_password = request.form.get("confirm_password")

    # Validate current password
    if not current_user.check_password(current_password):
        flash("Current password is incorrect")
        return redirect(url_for("profile"))

    # Validate new password
    if not new_password or len(new_password) < 4:
        flash("New password must be at least 4 characters long")
        return redirect(url_for("profile"))

    # Confirm passwords match
    if new_password != confirm_password:
        flash("New passwords do not match")
        return redirect(url_for("profile"))

    # Update password
    current_user.set_password(new_password)
    db.session.commit()

    flash("Password updated successfully")
    return redirect(url_for("profile"))


@app.route("/file_upload", methods=["GET", "POST"])
@login_required
def file_upload():
    """File upload page"""
    if request.method == "POST" and "files" in request.files:
        files = request.files.getlist("files")
        
        if not files or files[0].filename == "":
            flash("No files selected")
            return redirect(url_for("file_upload"))
        
        uploaded_count = 0
        
        for file in files:
            if file.filename:
                try:
                    # Save the file
                    secure_name = secure_filename(file.filename)
                    uploads_dir = os.path.join(app.config["BASE_DIR"], str(current_user.id), "uploads")
                    os.makedirs(uploads_dir, exist_ok=True)
                    file_path = os.path.join(uploads_dir, secure_name)
                    file.save(file_path)
                    
                    # Create database record
                    new_document = Document(
                        filename=secure_name,
                        file_path=file_path,
                        content_type=file.content_type,
                        user_id=current_user.id,
                        is_indexed=False
                    )
                    db.session.add(new_document)
                    db.session.commit()
                    
                    # Index the document (in a production app, this should be done in a background task)
                    index_document(new_document.id)
                    
                    uploaded_count += 1
                except Exception as e:
                    app.logger.error(f"Error uploading file {file.filename}: {str(e)}")
                    flash(f"Error uploading file {file.filename}: {str(e)}")
        
        if uploaded_count > 0:
            flash(f"Successfully uploaded {uploaded_count} file(s). Indexing in progress...")
        
        return redirect(url_for("file_upload"))
    
    # Get user's documents from database
    documents = Document.query.filter_by(user_id=current_user.id).order_by(Document.upload_date.desc()).all()
    
    return render_template("file_upload.html", documents=documents)


@app.route("/chat")
@login_required
def chat():
    """Unified chat interface for both regular chat and document chat"""
    # Get user's chat history
    chat_history_path = os.path.join(
        app.config["BASE_DIR"], str(current_user.id), "chat_history"
    )
    os.makedirs(chat_history_path, exist_ok=True)

    # Get list of conversations
    conversations = []
    try:
        for filename in os.listdir(chat_history_path):
            if filename.endswith(".json"):
                # Get the first message as title or use the filename
                filepath = os.path.join(chat_history_path, filename)
                with open(filepath, "r") as f:
                    try:
                        data = json.load(f)
                        title = (
                            data[0]["content"][:30] + "..."
                            if len(data[0]["content"]) > 30
                            else data[0]["content"]
                        )
                        timestamp = datetime.fromtimestamp(os.path.getctime(filepath))
                        conversations.append(
                            {
                                "id": os.path.splitext(filename)[0],
                                "title": title,
                                "timestamp": timestamp.strftime("%Y-%m-%d %H:%M:%S"),
                                "type": "chat",  # Mark as regular chat
                            }
                        )
                    except:
                        pass
    except:
        # Directory might not exist yet
        pass

    conversations.sort(key=lambda x: x["timestamp"], reverse=True)

    # Get user's indexed documents
    documents = Document.query.filter_by(user_id=current_user.id, is_indexed=True).all()

    return render_template(
        "unified_chat.html", conversations=conversations, documents=documents
    )


@app.route("/chat/<conversation_id>")
@login_required
def view_conversation(conversation_id):
    """View a specific conversation"""
    chat_history_path = os.path.join(
        app.config["BASE_DIR"], str(current_user.id), "chat_history"
    )
    os.makedirs(chat_history_path, exist_ok=True)

    # Get list of conversations
    conversations = []
    try:
        for filename in os.listdir(chat_history_path):
            if filename.endswith(".json"):
                filepath = os.path.join(chat_history_path, filename)
                with open(filepath, "r") as f:
                    try:
                        data = json.load(f)
                        title = (
                            data[0]["content"][:30] + "..."
                            if len(data[0]["content"]) > 30
                            else data[0]["content"]
                        )
                        timestamp = datetime.fromtimestamp(os.path.getctime(filepath))
                        conversations.append(
                            {
                                "id": os.path.splitext(filename)[0],
                                "title": title,
                                "timestamp": timestamp.strftime("%Y-%m-%d %H:%M:%S"),
                                "type": "chat",  # Mark as regular chat
                            }
                        )
                    except:
                        pass
    except:
        # Directory might not exist
        pass

    conversations.sort(key=lambda x: x["timestamp"], reverse=True)

    # Get conversation messages
    conversation_file = os.path.join(chat_history_path, f"{conversation_id}.json")
    messages = []

    if os.path.exists(conversation_file):
        with open(conversation_file, "r") as f:
            try:
                messages = json.load(f)
            except:
                # Invalid file
                pass

    # Get user's indexed documents
    documents = Document.query.filter_by(user_id=current_user.id, is_indexed=True).all()

    return render_template(
        "unified_chat.html",
        conversations=conversations,
        conversation_id=conversation_id,
        messages=messages,
        documents=documents,
    )


@app.route("/api/chat/new", methods=["POST"])
@login_required
def new_chat():
    """Start a new chat conversation"""
    try:
        message = request.json.get("message", "")

        if not message:
            return jsonify({"error": "No message provided"})

        # Calculate tokens
        user_tokens = count_tokens(message)

        # Generate a new conversation ID
        conversation_id = str(uuid.uuid4())

        # Get response from Azure OpenAI
        system_message = "You are a helpful AI assistant that can answer questions and provide help with various topics."

        # Create initial message list with just the user message
        messages = [{"role": "user", "content": message}]

        # Get AI response
        response = get_azure_openai_response(messages, system_message)

        # Calculate response tokens
        assistant_tokens = count_tokens(response)

        # Create messages with token counts
        messages = [
            {"role": "user", "content": message, "tokens": user_tokens},
            {"role": "assistant", "content": response, "tokens": assistant_tokens},
        ]

        # Create conversation data
        chat_data = {
            "id": conversation_id,
            "title": message[:30] + ("..." if len(message) > 30 else ""),
            "messages": messages,
            "created": datetime.now().timestamp(),
            "updated": datetime.now().timestamp(),
            "total_tokens": user_tokens + assistant_tokens,
        }

        # Save conversation
        save_conversation(conversation_id, messages)

        return jsonify(
            {
                "conversation_id": conversation_id,
                "messages": messages,
                "total_tokens": user_tokens + assistant_tokens,
            }
        )
    except Exception as e:
        app.logger.error(f"Error in new_chat: {str(e)}")
        return jsonify({"error": f"Server error: {str(e)}"})


@app.route("/api/chat/<conversation_id>/message", methods=["POST"])
@login_required
def add_message(conversation_id):
    """Add a message to an existing conversation"""
    try:
        message = request.json.get("message", "")

        if not message:
            return jsonify({"error": "No message provided"})

        # Calculate tokens
        user_tokens = count_tokens(message)

        # Load existing conversation
        chat_history_path = os.path.join(
            app.config["BASE_DIR"], str(current_user.id), "chat_history"
        )
        conversation_file = os.path.join(chat_history_path, f"{conversation_id}.json")

        if not os.path.exists(conversation_file):
            return jsonify({"error": "Conversation not found"})

        with open(conversation_file, "r") as f:
            try:
                messages = json.load(f)
            except:
                messages = []

        # Add new message with token count
        messages.append({"role": "user", "content": message, "tokens": user_tokens})

        # Get response from Azure OpenAI
        # Format messages for the API (skip token counts)
        api_messages = [{"role": m["role"], "content": m["content"]} for m in messages]

        # Get AI response
        response = get_azure_openai_response(api_messages)

        # Calculate response tokens
        assistant_tokens = count_tokens(response)

        # Add response to messages with token count
        messages.append(
            {"role": "assistant", "content": response, "tokens": assistant_tokens}
        )

        # Calculate total tokens for this conversation
        total_tokens = sum(m.get("tokens", 0) for m in messages)

        # Save updated conversation
        save_conversation(conversation_id, messages)

        return jsonify(
            {
                "conversation_id": conversation_id,
                "messages": messages,
                "total_tokens": total_tokens,
            }
        )
    except Exception as e:
        app.logger.error(f"Error in add_message: {str(e)}")
        return jsonify({"error": f"Server error: {str(e)}"})


@app.route("/api/chat/<conversation_id>/delete", methods=["POST"])
@login_required
def delete_conversation(conversation_id):
    """Delete a conversation"""
    chat_history_path = os.path.join(
        app.config["BASE_DIR"], str(current_user.id), "chat_history"
    )
    conversation_file = os.path.join(chat_history_path, f"{conversation_id}.json")

    if os.path.exists(conversation_file):
        os.remove(conversation_file)

    return jsonify({"success": True})


def save_conversation(conversation_id, messages):
    """Save conversation to local storage"""
    chat_history_path = os.path.join(
        app.config["BASE_DIR"], str(current_user.id), "chat_history"
    )
    os.makedirs(chat_history_path, exist_ok=True)

    conversation_file = os.path.join(chat_history_path, f"{conversation_id}.json")

    with open(conversation_file, "w") as f:
        json.dump(messages, f)

    return True


# Helper function for file uploads
def allowed_file(filename):
    # Allow all files for now
    return True


if __name__ == "__main__":
    # Create database tables if they don't exist
    with app.app_context():
        db.create_all()

        # Create admin user if there are no users
        if not User.query.first():
            admin = User(
                username="admin",
                email="admin@example.com",
                is_admin=True,
                is_approved=True,
            )
            admin.set_password("admin123")
            db.session.add(admin)

            user = User(username="user", email="user@example.com", is_approved=True)
            user.set_password("user123")
            db.session.add(user)

            db.session.commit()

            # Create user directories
            ensure_user_dirs(admin.id)
            ensure_user_dirs(user.id)

    app.run(debug=True)
